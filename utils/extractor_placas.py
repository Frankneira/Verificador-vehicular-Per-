"""
extractor_placas.py -- Extrae placas peruanas de archivos Excel, Word y PDF.

Formatos validos de placa peruana:
  ABC-123  |  ABC123  |  B1N-553  |  D6R-678
  Las placas peruanas SIEMPRE comienzan con una LETRA.

Retorna lista de placas unicas detectadas, en mayusculas, sin guiones.
"""
import re
import io
from pathlib import Path

# Placa peruana: SIEMPRE empieza con letra, luego mezcla alfanum, termina en 2-3 digitos
# Ejemplos: ABC123, B1N553, D6R678, A8R215, ABC-123
PLACA_RE = re.compile(
    r'\b([A-Z][A-Z0-9]{1,2}[-\s]?[A-Z0-9]{0,2}[-\s]?[0-9]{2,3})\b',
    re.IGNORECASE
)

# Patron estricto post-limpieza: 5-6 chars alfanum
PLACA_ESTRICTA = re.compile(r'^[A-Z][A-Z0-9]{1,2}[0-9A-Z]{2,3}$')

# Palabras a ignorar
IGNORAR = {
    "DNI", "RUC", "NUM", "NRO", "COD", "REF", "TEL",
    "FAX", "CEL", "ANO", "AYO", "URL", "PDF", "XML",
}


def _limpiar(placa_raw: str) -> str:
    """Elimina espacios y guiones, convierte a mayusculas."""
    return re.sub(r'[\s\-]', '', placa_raw.strip().upper())


def _es_placa_valida(token: str) -> bool:
    """Valida que el token tenga forma de placa peruana."""
    limpia = _limpiar(token)
    if len(limpia) < 5 or len(limpia) > 6:
        return False
    if limpia in IGNORAR:
        return False
    if not limpia[0].isalpha():
        return False
    letras  = sum(1 for c in limpia if c.isalpha())
    digitos = sum(1 for c in limpia if c.isdigit())
    if letras < 2 or digitos < 2:
        return False
    return bool(PLACA_ESTRICTA.match(limpia))


def _normalizar_celda(valor) -> str:
    """Convierte cualquier valor de celda a string limpio."""
    if valor is None:
        return ""
    return str(valor).strip()


def extraer_de_texto(texto: str) -> list:
    """Extrae placas de un string de texto plano."""
    candidatos = PLACA_RE.findall(texto)
    placas = []
    vistas = set()
    for c in candidatos:
        limpia = _limpiar(c)
        if _es_placa_valida(limpia) and limpia not in vistas:
            placas.append(limpia)
            vistas.add(limpia)
    return placas


def extraer_de_excel(contenido: bytes) -> list:
    """
    Extrae placas de un archivo Excel (.xlsx / .xls).
    Estrategia: procesa cada celda individualmente para evitar
    que el numero de fila se concatene con la placa.
    """
    vistas = set()
    placas = []

    def _procesar_valor(val):
        s = _normalizar_celda(val)
        if not s:
            return
        # Intentar como placa exacta primero
        limpia = _limpiar(s)
        if _es_placa_valida(limpia) and limpia not in vistas:
            placas.append(limpia)
            vistas.add(limpia)
            return
        # Si la celda tiene texto largo, buscar placas dentro
        if len(s) > 7:
            for c in PLACA_RE.findall(s):
                limpia2 = _limpiar(c)
                if _es_placa_valida(limpia2) and limpia2 not in vistas:
                    placas.append(limpia2)
                    vistas.add(limpia2)

    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for celda in row:
                    _procesar_valor(celda)
        return placas
    except Exception:
        pass

    try:
        import xlrd
        book = xlrd.open_workbook(file_contents=contenido)
        for sheet in book.sheets():
            for r in range(sheet.nrows):
                for c in range(sheet.ncols):
                    _procesar_valor(sheet.cell_value(r, c))
        return placas
    except Exception as e:
        raise ValueError("No se pudo leer el Excel: {}".format(e))


def extraer_de_word(contenido: bytes) -> list:
    """Extrae placas de un archivo Word (.docx)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(contenido))
        partes = []
        for parrafo in doc.paragraphs:
            partes.append(parrafo.text)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    partes.append(celda.text)
        return extraer_de_texto(" ".join(partes))
    except Exception as e:
        raise ValueError("No se pudo leer el Word: {}".format(e))


def extraer_de_pdf(contenido: bytes) -> list:
    """Extrae placas de un archivo PDF."""
    try:
        import pdfplumber
        texto_total = ""
        with pdfplumber.open(io.BytesIO(contenido)) as pdf:
            for pagina in pdf.pages:
                t = pagina.extract_text()
                if t:
                    texto_total += " " + t
        return extraer_de_texto(texto_total)
    except Exception as e:
        raise ValueError("No se pudo leer el PDF: {}".format(e))


def extraer_placas(nombre_archivo: str, contenido: bytes) -> list:
    """
    Entrada principal. Detecta el tipo de archivo y extrae placas.

    Args:
        nombre_archivo : Nombre del archivo (para detectar extension).
        contenido      : Bytes del archivo.

    Returns:
        Lista de placas unicas en mayusculas sin guiones. Ej: ["D6R678", "B1N553"]

    Raises:
        ValueError si el formato no es soportado.
    """
    ext = Path(nombre_archivo).suffix.lower()

    if ext in (".xlsx", ".xls"):
        return extraer_de_excel(contenido)
    elif ext in (".docx", ".doc"):
        return extraer_de_word(contenido)
    elif ext == ".pdf":
        return extraer_de_pdf(contenido)
    elif ext in (".txt", ".csv"):
        return extraer_de_texto(contenido.decode("utf-8", errors="replace"))
    else:
        raise ValueError(
            "Formato no soportado: '{}'. "
            "Usa Excel (.xlsx), Word (.docx), PDF (.pdf) o TXT/CSV.".format(ext)
        )
